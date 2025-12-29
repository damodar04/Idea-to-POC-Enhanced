"""Research Document Generator - Creates comprehensive documents with research insights"""

import logging
import os
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from langchain_openai import AzureChatOpenAI, ChatOpenAI
from langchain_core.prompts import PromptTemplate

logger = logging.getLogger(__name__)

class ResearchDocumentGenerator:
    """Generates comprehensive documents with research insights"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.llm = None
        self._init_llm()

    def _init_llm(self):
        """Initialize the LLM for content generation"""
        try:
            # Use existing Azure OpenAI configuration from environment
            gpt_4o_api_key = os.getenv("GPT_4O_API_KEY")
            azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
            
            if not gpt_4o_api_key or not azure_endpoint:
                # Try to load from .env if not found
                from dotenv import load_dotenv
                load_dotenv()
                gpt_4o_api_key = os.getenv("GPT_4O_API_KEY")
                azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
                deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
            
            # Try Azure OpenAI first
            if gpt_4o_api_key and azure_endpoint and not gpt_4o_api_key.startswith("your_"):
                # Clean endpoint if needed
                if "openai/deployments" in azure_endpoint:
                    from urllib.parse import urlparse
                    parsed = urlparse(azure_endpoint)
                    azure_endpoint = f"{parsed.scheme}://{parsed.netloc}/"

                self.llm = AzureChatOpenAI(
                    api_key=gpt_4o_api_key,
                    azure_endpoint=azure_endpoint,
                    api_version="2024-02-01",
                    azure_deployment="gpt-4o",
                    temperature=0.7
                )
                self.logger.info("Azure OpenAI GPT-4o initialized for document generation")
            # Fallback to DeepSeek
            elif deepseek_api_key and not deepseek_api_key.startswith("your_"):
                self.llm = ChatOpenAI(
                    api_key=deepseek_api_key,
                    base_url="https://api.deepseek.com",
                    model="deepseek-chat",
                    temperature=0.7
                )
                self.logger.info("DeepSeek initialized for document generation")
            else:
                self.logger.warning("No AI API credentials found. Document generation will use raw data only.")

        except Exception as e:
            self.logger.error(f"LLM initialization failed: {str(e)}")
    
    def generate_comprehensive_document(self, idea_data: Dict) -> Optional[bytes]:
        """
        Generate a comprehensive document with research insights
        
        Args:
            idea_data: Complete idea data including research and answers
            
        Returns:
            DOCX document bytes or None if generation fails
        """
        try:
            self.logger.info("Generating comprehensive research document")
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(level=1)
            title_run = title.add_run(f"Proof of Concept Proposal: {idea_data.get('idea_title', 'Untitled Idea')}")
            title_run.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add company info
            company_para = doc.add_paragraph()
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            company_run = company_para.add_run(f"Prepared for: {idea_data.get('company_name', 'Unknown Company')}")
            company_run.italic = True
            
            doc.add_paragraph()  # Add spacing
            
            # Generate AI content if LLM is available
            ai_content = {}
            if self.llm:
                self.logger.info("Synthesizing document sections with AI...")
                ai_content = self._generate_all_sections_ai(idea_data)
            
            # Add Executive Summary
            self._add_section_with_fallback(doc, "Executive Summary", ai_content.get('executive_summary'), idea_data, self._add_executive_summary_fallback)
            
            # Add Problem Statement
            self._add_section_with_fallback(doc, "Problem Statement", ai_content.get('problem_statement'), idea_data, self._add_problem_statement_fallback)

            # Add Proposed Solution
            self._add_section_with_fallback(doc, "Proposed Solution", ai_content.get('proposed_solution'), idea_data, self._add_proposed_solution_fallback)
            
            # Add Strategic Alignment
            self._add_section_with_fallback(doc, "Strategic Alignment", ai_content.get('strategic_alignment'), idea_data, self._add_strategic_alignment_fallback)

            # Add Market Analysis
            self._add_section_with_fallback(doc, "Market Analysis", ai_content.get('market_analysis'), idea_data, self._add_market_research_fallback)
            
            # Add Implementation Plan
            self._add_section_with_fallback(doc, "Implementation Plan", ai_content.get('implementation_plan'), idea_data, self._add_resource_estimation_fallback)
            
            # Add Success Metrics
            self._add_section_with_fallback(doc, "Success Metrics & ROI", ai_content.get('success_metrics'), idea_data, self._add_success_metrics_fallback)
            
            # Add Conclusion
            self._add_section_with_fallback(doc, "Conclusion", ai_content.get('conclusion'), idea_data, self._add_conclusion_fallback)
            
            # Save to bytes
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            self.logger.info("Comprehensive research document generated successfully")
            return doc_bytes.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error generating comprehensive document: {str(e)}")
            import traceback
            self.logger.error(traceback.format_exc())
            return None

    def _generate_all_sections_ai(self, idea_data: Dict) -> Dict:
        """Generate all document sections using AI"""
        sections = {}
        
        # Prepare context string
        context = self._prepare_context(idea_data)
        
        # We can generate sections in parallel or sequence. Sequence for now to keep it simple.
        sections['executive_summary'] = self._generate_section_ai("Executive Summary", context, "Write a compelling executive summary for this POC proposal. Focus on the value proposition and key outcomes. Keep it high-level and brief.")
        sections['problem_statement'] = self._generate_section_ai("Problem Statement", context, "Describe the current business challenges and market gaps this idea addresses. Focus ONLY on the problem. Do not discuss the solution yet.")
        sections['proposed_solution'] = self._generate_section_ai("Proposed Solution", context, "Detail the proposed solution, including technical approach and key features. Focus ONLY on the solution details. Do not repeat the problem.")
        sections['strategic_alignment'] = self._generate_section_ai("Strategic Alignment", context, "Explain how this idea aligns with the company's current initiatives, goals, and financial context.")
        sections['market_analysis'] = self._generate_section_ai("Market Analysis", context, "Analyze the market landscape, including competitors, trends, and existing solutions. Focus on external market factors.")
        sections['implementation_plan'] = self._generate_section_ai("Implementation Plan", context, "Outline the implementation strategy, timeline, required resources, and risk mitigation strategies. Use bullet points for steps.")
        sections['success_metrics'] = self._generate_section_ai("Success Metrics", context, "Define the key performance indicators (KPIs) and projected ROI. Use bullet points for metrics.")
        sections['conclusion'] = self._generate_section_ai("Conclusion", context, "Summarize the proposal and recommend clear next steps. Be very concise.")
        
        return sections

    def _prepare_context(self, idea_data: Dict) -> str:
        """Prepare a text representation of all available data for the AI"""
        context = []
        
        # Handle title variations
        title = idea_data.get('idea_title') or idea_data.get('title') or 'Untitled Idea'
        context.append(f"Idea Title: {title}")
        
        context.append(f"Company: {idea_data.get('company_name')}")
        context.append(f"Description: {idea_data.get('idea_description')}")
        
        # Company Research
        cr = idea_data.get('company_research', {})
        if cr:
            context.append("\nCOMPANY RESEARCH:")
            context.append(f"Overview: {cr.get('what_company_does')}")
            context.append(f"Financials: {cr.get('financials')}")
            context.append(f"Goals: {cr.get('current_initiatives_and_goals')}")
            context.append(f"Challenges: {cr.get('challenges')}")
            
        # Market Research
        mr = idea_data.get('market_research') or idea_data.get('idea_research') or {}
        if mr:
            context.append("\nMARKET RESEARCH:")
            context.append(f"Overview: {mr.get('market_overview')}")
            context.append(f"Existing Solutions: {mr.get('existing_solutions')}")
            context.append(f"Competitors: {mr.get('competitors')}")
            context.append(f"Insights: {mr.get('useful_insights')}")
            
        # Development Answers (User Input)
        answers = idea_data.get('development_answers', {})
        questions = idea_data.get('development_questions', [])
        if answers:
            context.append("\nDETAILED DEVELOPMENT PLAN (User Answers):")
            for q in questions:
                key = q.get('key')
                if key in answers:
                    context.append(f"Q ({q.get('section')}): {q.get('question')}")
                    context.append(f"A: {answers[key]}")
                    
        # Resource Estimation
        re = idea_data.get('resource_estimation', {})
        if re:
             context.append("\nRESOURCE ESTIMATION:")
             context.append(str(re))

        # ROI Analysis
        roi = idea_data.get('roi_analysis', {})
        if roi:
            context.append("\nROI ANALYSIS:")
            context.append(str(roi))
            
        return "\n".join(context)

    def _generate_section_ai(self, section_name: str, context: str, instruction: str) -> str:
        """Generate a single section using AI"""
        try:
            prompt = PromptTemplate(
                input_variables=["section_name", "instruction", "context"],
                template="""You are an expert business consultant writing a Proof of Concept (POC) proposal.
                
Context Information:
{context}

Task: Write the '{section_name}' section of the document.
Instruction: {instruction}

Requirements:
- Write in a professional, persuasive business tone.
- Use bullet points (start lines with "- ") for lists, key features, or distinct points to improve readability and avoid walls of text.
- Do NOT use other markdown formatting like **bold** or # headers.
- Avoid repeating information. Be concise and direct.
- Synthesize the information from the context.
- Length: Adequate to cover the topic but concise.

Content:"""
            )
            
            return self.llm.invoke(prompt.format(section_name=section_name, instruction=instruction, context=context[:50000])).content
            
        except Exception as e:
            self.logger.error(f"Failed to generate section {section_name}: {e}")
            return ""

    def _add_section_with_fallback(self, doc: Document, title: str, ai_content: Optional[str], idea_data: Dict, fallback_method):
        """Add a section to the document, using AI content if available, otherwise fallback"""
        doc.add_heading(title, level=2)
        
        if ai_content and len(ai_content.strip()) > 50:
            # Split by newlines to handle potential bullet points
            lines = ai_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('- ') or line.startswith('* '):
                    # It's a bullet point
                    clean_line = line[2:].strip()
                    doc.add_paragraph(clean_line, style='List Bullet')
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
        else:
            fallback_method(doc, idea_data)

    def _add_executive_summary_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for executive summary"""
        summary_para = doc.add_paragraph()
        summary_para.add_run("This document presents a comprehensive business idea proposal based on extensive market research, company analysis, and financial viability assessment. ")
        summary_para.add_run(f"The idea '{idea_data.get('idea_title', 'Untitled Idea')}' targets {idea_data.get('company_name', 'the selected company')} and addresses specific market opportunities identified through research.")
    
    def _add_problem_statement_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for problem statement"""
        company_research = idea_data.get('company_research', {})
        if company_research.get('challenges'):
            doc.add_paragraph("The following business challenges have been identified:")
            for challenge in company_research['challenges'][:5]:
                doc.add_paragraph(self._clean_text(challenge), style='List Bullet')
        else:
            doc.add_paragraph("No specific problem statement data available.")

    def _add_proposed_solution_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for proposed solution"""
        doc.add_paragraph(idea_data.get('idea_description', ''))
        
        # Add development answers if available
        answers = idea_data.get('development_answers', {})
        if answers:
            doc.add_heading("Implementation Details", level=3)
            for key, answer in answers.items():
                if answer:
                    doc.add_paragraph(self._clean_text(answer))

    def _add_strategic_alignment_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for strategic alignment"""
        company_research = idea_data.get('company_research', {})
        if company_research.get('current_initiatives_and_goals'):
            doc.add_paragraph("This proposal aligns with the following company initiatives:")
            for init in company_research['current_initiatives_and_goals'][:5]:
                doc.add_paragraph(self._clean_text(init), style='List Bullet')

    def _add_market_research_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for market research (reusing old logic)"""
        self._add_market_research(doc, idea_data) # Use the existing method structure

    def _add_resource_estimation_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for resource estimation (reusing old logic)"""
        self._add_resource_estimation(doc, idea_data)

    def _add_success_metrics_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for success metrics"""
        resource_est = idea_data.get('resource_estimation', {})
        if not resource_est:
             resource_est = idea_data.get('research_data', {}).get('resource_estimation', {})

        if resource_est.get('success_metrics'):
            for metric in resource_est['success_metrics']:
                metric_name = self._clean_text(metric.get('metric', 'Success Metric'))
                target = self._clean_text(metric.get('target_value', metric.get('target', 'N/A')))
                doc.add_paragraph(f"{metric_name}: {target}", style='List Bullet')
        
        roi = idea_data.get('roi_analysis', {})
        if roi:
            doc.add_paragraph(f"Estimated ROI Score: {roi.get('roi_score', 'N/A')}%")

    def _add_conclusion_fallback(self, doc: Document, idea_data: Dict):
        """Fallback for conclusion"""
        doc.add_paragraph("This comprehensive analysis demonstrates the viability and potential of the proposed business idea.")

    # Keep existing helper methods
    def _add_market_research(self, doc: Document, idea_data: Dict):
        """Add market research section (Legacy/Fallback)"""
        idea_research = idea_data.get('idea_research')
        if not idea_research or not idea_research.get('success'):
            doc.add_paragraph("No market research data available.")
            return
        
        # Companies implementing this idea
        if idea_research.get('who_is_implementing'):
            doc.add_heading("Companies Implementing This Idea", level=3)
            for implementer in idea_research['who_is_implementing'][:8]:
                name = self._clean_text(implementer.get('name', 'Company'))
                description = self._clean_text(implementer.get('description', ''))
                
                para = doc.add_paragraph(style='List Bullet')
                name_run = para.add_run(f"{name}: ")
                name_run.bold = True
                para.add_run(description)
        
        # Pros and Cons
        pros_cons = idea_research.get('pros_and_cons', {})
        
        if pros_cons.get('pros'):
            doc.add_heading("Implementation Benefits", level=3)
            for pro in pros_cons['pros'][:8]:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(self._clean_text(pro))
        
        if pros_cons.get('cons'):
            doc.add_heading("Implementation Challenges", level=3)
            for con in pros_cons['cons'][:8]:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(self._clean_text(con))
        
        # Market Insights
        if idea_research.get('useful_insights'):
            doc.add_heading("Key Market Insights", level=3)
            for insight in idea_research['useful_insights'][:10]:
                insight_type = self._clean_text(insight.get('type', 'Insight'))
                insight_text = self._clean_text(insight.get('insight', ''))
                details = self._clean_text(insight.get('details', ''))
                
                para = doc.add_paragraph(style='List Bullet')
                type_run = para.add_run(f"{insight_type}: ")
                type_run.bold = True
                para.add_run(insight_text)
                if details:
                    para.add_run(f" ({details})")

    def _add_resource_estimation(self, doc: Document, idea_data: Dict):
        """Add resource estimation section (Legacy/Fallback)"""
        resource_est = idea_data.get('resource_estimation')
        if not resource_est:
            resource_est = idea_data.get('research_data', {}).get('resource_estimation')

        if not resource_est or not resource_est.get('success'):
            doc.add_paragraph("No resource estimation data available.")
            return
        
        # Team Resources
        if resource_est.get('team_resources'):
            doc.add_heading("Team Resources Required", level=3)
            for resource in resource_est['team_resources']:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(self._clean_text(resource.get('description', 'Team Member')))
        
        # Implementation Timeline
        if resource_est.get('timeline'):
            doc.add_heading("Implementation Timeline", level=3)
            for phase in resource_est['timeline']:
                phase_name = self._clean_text(phase.get('phase', 'Phase'))
                para = doc.add_paragraph(style='List Bullet')
                phase_run = para.add_run(f"{phase_name}")
                phase_run.bold = True
                if phase.get('duration'):
                    para.add_run(f" - Duration: {self._clean_text(phase['duration'])}")
        
        # Technical Infrastructure
        if resource_est.get('technical_infrastructure'):
            doc.add_heading("Technical Infrastructure", level=3)
            for item in resource_est['technical_infrastructure']:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(self._clean_text(item))
        
        # Risk Assessment
        if resource_est.get('risks'):
            doc.add_heading("Risk Assessment", level=3)
            for risk in resource_est['risks']:
                risk_desc = self._clean_text(risk.get('risk', 'Risk identified'))
                impact = self._clean_text(risk.get('impact_level', risk.get('impact', 'Medium')))
                mitigation = self._clean_text(risk.get('mitigation_strategy', risk.get('mitigation', 'N/A')))
                
                para = doc.add_paragraph(style='List Bullet')
                risk_run = para.add_run(f"{risk_desc}")
                risk_run.bold = True
                para.add_run(f" (Impact: {impact})")
                doc.add_paragraph(f"Mitigation: {mitigation}", style='List Bullet 2')

    def _clean_text(self, text: str) -> str:
        """Clean text by removing markdown formatting"""
        if not isinstance(text, str):
            return str(text)
        
        # Remove bold markers
        text = text.replace('**', '').replace('__', '')
        # Remove italic markers (simple check)
        if text.startswith('*') and text.endswith('*'):
            text = text[1:-1]
        elif text.startswith('_') and text.endswith('_'):
            text = text[1:-1]
            
        return text.strip()

# Create global instance
research_document_generator = ResearchDocumentGenerator()
